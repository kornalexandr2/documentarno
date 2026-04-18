import logging
import os
import tempfile
from typing import Callable, Optional

import fitz  # PyMuPDF
import torch
from docx import Document as DocxDocument
from markdownify import markdownify as md
from paddleocr import PPStructure

logger = logging.getLogger(__name__)

_OCR_ENGINES: dict[str, PPStructure] = {}
_CPU_ONLY = os.getenv("OCR_FORCE_CPU", "0").lower() in {"1", "true", "yes"}


def _create_ocr_engine(use_gpu: bool) -> PPStructure:
    engine = PPStructure(
        lang="en",
        show_log=False,
        use_gpu=use_gpu,
        layout=True,
        table=True,
        recovery=True,
    )
    logger.info("PaddleOCR engine initialized successfully (GPU: %s).", use_gpu)
    return engine


def _get_ocr_engine(prefer_gpu: bool = True) -> PPStructure:
    allow_gpu = prefer_gpu and not _CPU_ONLY and torch.cuda.is_available()
    engine_key = "gpu" if allow_gpu else "cpu"

    if engine_key in _OCR_ENGINES:
        return _OCR_ENGINES[engine_key]

    try:
        engine = _create_ocr_engine(use_gpu=allow_gpu)
        _OCR_ENGINES[engine_key] = engine
        return engine
    except Exception as exc:
        if allow_gpu:
            logger.warning("GPU OCR initialization failed, falling back to CPU: %s", exc)
            return _get_ocr_engine(prefer_gpu=False)
        raise RuntimeError(f"Failed to initialize OCR engine: {exc}") from exc


def _is_gpu_runtime_error(exc: Exception) -> bool:
    message = str(exc).lower()
    return "cudnn" in message or "preconditionnotmet" in message or "cuda" in message


def process_pdf_to_markdown(file_path: str, progress_callback: Optional[Callable[[int, int], None]] = None) -> str:
    """
    Extract text and tables from a PDF using PaddleOCR PPStructure.
    Falls back to CPU if GPU runtime dependencies are missing.
    """
    doc = None
    active_engine = _get_ocr_engine(prefer_gpu=True)
    using_gpu = active_engine is _OCR_ENGINES.get("gpu")

    try:
        doc = fitz.open(file_path)
        markdown_content = []
        temp_dir = tempfile.gettempdir()
        total_pages = len(doc)

        for page_num in range(total_pages):
            if progress_callback:
                progress_callback(page_num + 1, total_pages)

            logger.info("--- [OCR PROGRESS] Processing page %s/%s ---", page_num + 1, total_pages)
            page = doc.load_page(page_num)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_bytes = pix.tobytes("png")

            temp_img_path = os.path.join(temp_dir, f"doc_page_{page_num}_{os.getpid()}.png")
            with open(temp_img_path, "wb") as temp_file:
                temp_file.write(img_bytes)

            try:
                try:
                    result = active_engine(temp_img_path)
                except Exception as exc:
                    if using_gpu and _is_gpu_runtime_error(exc):
                        logger.warning("GPU OCR runtime failed on page %s, retrying on CPU: %s", page_num + 1, exc)
                        active_engine = _get_ocr_engine(prefer_gpu=False)
                        using_gpu = False
                        result = active_engine(temp_img_path)
                    else:
                        raise

                markdown_content.append(f"\n## Страница {page_num + 1}\n")

                for region in result:
                    region_type = region["type"]
                    res = region["res"]

                    if region_type == "Table":
                        html_table = res.get("html", "")
                        if html_table:
                            md_table = md(html_table, strip=["html", "body", "head"])
                            markdown_content.append(md_table.strip())
                    elif region_type in ["Figure", "Equation"]:
                        markdown_content.append("\n*[Изображение/График]*\n")
                    else:
                        text_lines = [line["text"] for line in res]
                        text_block = " ".join(text_lines)
                        if region_type == "Title":
                            markdown_content.append(f"\n### {text_block}\n")
                        else:
                            markdown_content.append(text_block)
            finally:
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)

        return "\n".join(markdown_content)
    finally:
        if doc is not None:
            doc.close()


def process_docx_to_markdown(file_path: str) -> str:
    """
    Extract text and simple tables from a DOCX file and convert it to Markdown.
    """
    try:
        logger.info("--- [DOCX PROGRESS] Extracting text from %s ---", os.path.basename(file_path))
        doc = DocxDocument(file_path)
        markdown_content = []

        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith("Heading 1"):
                    markdown_content.append(f"# {para.text}")
                elif para.style.name.startswith("Heading 2"):
                    markdown_content.append(f"## {para.text}")
                elif para.style.name.startswith("Heading 3"):
                    markdown_content.append(f"### {para.text}")
                else:
                    markdown_content.append(para.text)

        for table in doc.tables:
            markdown_content.append("\n")
            if len(table.rows) > 0:
                header_row = table.rows[0]
                cols_count = len(header_row.cells)
                for i, row in enumerate(table.rows):
                    cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                    markdown_content.append("| " + " | ".join(cells) + " |")
                    if i == 0:
                        markdown_content.append("| " + " | ".join(["---"] * cols_count) + " |")
            markdown_content.append("\n")

        return "\n".join(markdown_content)
    except Exception as exc:
        logger.error("Error processing DOCX: %s", exc)
        raise RuntimeError(f"Failed to process DOCX: {exc}") from exc
